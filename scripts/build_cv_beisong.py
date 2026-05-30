from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase.pdfmetrics import stringWidth
from reportlab.pdfgen import canvas


OUT_DIR = Path("outputs/cv_beisong")
DOCX_PATH = OUT_DIR / "CV_Beisong_LYU_Apprentissage_IMT_Nord_Europe_2026.docx"
PDF_PATH = OUT_DIR / "CV_Beisong_LYU_Apprentissage_IMT_Nord_Europe_2026.pdf"

BLUE = RGBColor(31, 78, 121)
MUTED = RGBColor(90, 90, 90)
INK = RGBColor(30, 30, 30)


def set_cell_free_document_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.05)
    section.bottom_margin = Cm(1.05)
    section.left_margin = Cm(1.25)
    section.right_margin = Cm(1.25)
    section.header_distance = Cm(0.6)
    section.footer_distance = Cm(0.6)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(9.4)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(1.6)
    normal.paragraph_format.line_spacing = 1.02

    for style_name, size, color in [
        ("Heading 1", 9.6, BLUE),
        ("Heading 2", 9.2, BLUE),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(5)
        style.paragraph_format.space_after = Pt(1.5)
        style.paragraph_format.keep_with_next = True


def add_bottom_border(paragraph, color="D9E2F3") -> None:
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)


def paragraph(doc, text="", *, bold=False, italic=False, size=None, color=None, align=None, after=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    return p


def section_heading(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(1.2)
    r = p.add_run(text.upper())
    r.bold = True
    r.font.size = Pt(9.6)
    r.font.color.rgb = BLUE
    add_bottom_border(p)
    return p


def role_line(doc, left, right=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0.8)
    r = p.add_run(left)
    r.bold = True
    r.font.size = Pt(9.4)
    if right:
        r2 = p.add_run(f" | {right}")
        r2.font.size = Pt(9.4)
        r2.font.color.rgb = MUTED
    return p


def detail(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.18)
    p.paragraph_format.space_after = Pt(1.1)
    r = p.add_run(text)
    r.font.size = Pt(9.1)
    return p


def compact_list(doc, items):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.18)
    p.paragraph_format.space_after = Pt(1.0)
    for idx, item in enumerate(items):
        if idx:
            sep = p.add_run("  |  ")
            sep.font.color.rgb = MUTED
        r = p.add_run(item)
        r.font.size = Pt(9.05)
    return p


def build_cv():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    set_cell_free_document_defaults(doc)

    name = paragraph(doc, "LYU Beisong", bold=True, size=17.5, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    name.paragraph_format.space_before = Pt(0)

    subtitle = paragraph(
        doc,
        "Candidat admissible IMT Nord Europe | Recherche contrat d'apprentissage ingénieur 36 mois | Rentrée 2026",
        bold=True,
        size=10.2,
        color=INK,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=0.5,
    )
    subtitle.paragraph_format.keep_with_next = True

    contact = paragraph(
        doc,
        "Tours | Mobilité Lille, Douai et Hauts-de-France | 06 28 73 21 24 | benson201814@gmail.com | LinkedIn : à compléter | GitHub/Portfolio : à compléter",
        size=8.5,
        color=MUTED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=3,
    )
    add_bottom_border(contact, color="BFBFBF")

    section_heading(doc, "Profil")
    paragraph(
        doc,
        "Admissible à IMT Nord Europe pour une formation d'ingénieur par apprentissage, je recherche un contrat de 36 mois en informatique, télécommunications, réseaux ou data. Issu d'une CPGE MP2I-MPI, je combine une base solide en mathématiques, algorithmique et programmation avec un parcours international Chine - États-Unis - France. Trilingue chinois, anglais et français, je m'adapte rapidement aux environnements exigeants et multiculturels.",
        size=9.15,
        after=1.4,
    )

    section_heading(doc, "Formation")
    role_line(doc, "2024 - 2026 | CPGE MP2I-MPI - Lycée Descartes, Tours")
    detail(doc, "Mathématiques, informatique, algorithmique, physique. Compétences développées : rigueur scientifique, résolution de problèmes complexes, autonomie, forte capacité de travail.")
    detail(doc, "Responsabilités : délégué, membre du conseil d'administration.")

    role_line(doc, "2026 - 2029 | IMT Nord Europe", "admissible, contrat d'apprentissage à finaliser")
    detail(doc, "Formation d'ingénieur par apprentissage visée. Domaine cible : Informatique, Télécommunications et Réseaux.")

    role_line(doc, "2021 - 2024 | Baccalauréat général - Lycée Jean Lurçat")
    detail(doc, "Spécialités scientifiques. Vice-président du CVL, membre du conseil d'administration, délégué.")

    section_heading(doc, "Compétences techniques")
    compact_list(doc, ["Python", "C", "OCaml", "bases SQL", "algorithmique", "structures de données"])
    compact_list(doc, ["analyse fréquentielle", "traitement numérique du signal", "Git", "Linux", "Windows"])
    compact_list(doc, ["analyse", "synthèse", "résolution de problèmes", "apprentissage rapide", "communication interculturelle"])

    section_heading(doc, "Projets")
    role_line(doc, "TIPE - Analyse et synthèse du signal audio")
    detail(doc, "Implémentation d'algorithmes de synthèse sonore ; segmentation et analyse fréquentielle de signaux audio ; présentation d'une démarche scientifique structurée.")

    role_line(doc, "Développement de jeux indépendants")
    detail(doc, "Conception de prototypes interactifs ; programmation de mécaniques de jeu et de logique applicative ; tests et amélioration en autonomie.")

    role_line(doc, "Création de contenu audiovisuel")
    detail(doc, "Production de plus de 30 projets audiovisuels, 200 000 vues cumulées. Compétences : narration, montage, régularité et communication.")

    section_heading(doc, "Expériences et engagements")
    role_line(doc, "2023 | Assistant traducteur et support technique", "Projet UiT The Arctic University of Norway")
    detail(doc, "Création et gestion d'un site web ; traduction et interprétation chinois - anglais - français ; coordination et communication dans un contexte international.")

    role_line(doc, "2021 | Volontariat d'enseignement de l'anglais", "Tibet, Chine")
    detail(doc, "Enseignement auprès d'élèves en zone rurale ; animation de groupe et adaptation pédagogique.")

    role_line(doc, "2020 | Interprète bilingue", "Camp d'été, Portland, États-Unis")
    detail(doc, "Traduction chinois - anglais ; support logistique et médiation interculturelle.")

    section_heading(doc, "Langues et centres d'intérêt")
    compact_list(doc, ["Chinois : langue maternelle", "Anglais : courant, 4 ans de scolarité aux États-Unis", "Français : courant"])
    compact_list(doc, ["Piano", "course à pied", "jeux vidéo indépendants", "communication interculturelle"])

    doc.save(DOCX_PATH)
    print(DOCX_PATH.resolve())


def draw_wrapped(c, text, x, y, max_width, font="Helvetica", size=8.6, leading=9.6, color=colors.black):
    c.setFont(font, size)
    c.setFillColor(color)
    words = text.split()
    lines = []
    current = ""
    for word in words:
        candidate = word if not current else f"{current} {word}"
        if stringWidth(candidate, font, size) <= max_width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    for line in lines:
        c.drawString(x, y, line)
        y -= leading
    return y


def draw_section(c, title, x, y, width):
    y -= 7
    c.setFont("Helvetica-Bold", 8.8)
    c.setFillColor(colors.HexColor("#1F4E79"))
    c.drawString(x, y, title.upper())
    c.setStrokeColor(colors.HexColor("#D9E2F3"))
    c.setLineWidth(0.7)
    c.line(x, y - 2.5, x + width, y - 2.5)
    return y - 12


def draw_role(c, left, right, x, y, width):
    c.setFont("Helvetica-Bold", 8.5)
    c.setFillColor(colors.black)
    c.drawString(x, y, left)
    if right:
        c.setFont("Helvetica", 8.2)
        c.setFillColor(colors.HexColor("#555555"))
        right_text = f"| {right}"
        c.drawRightString(x + width, y, right_text)
    return y - 9.5


def build_pdf():
    c = canvas.Canvas(str(PDF_PATH), pagesize=A4)
    width, height = A4
    margin_x = 33
    y = height - 31
    usable = width - 2 * margin_x

    c.setTitle("CV Beisong LYU - Apprentissage IMT Nord Europe 2026")
    c.setAuthor("Beisong LYU")

    c.setFont("Helvetica-Bold", 17)
    c.setFillColor(colors.HexColor("#1F4E79"))
    c.drawCentredString(width / 2, y, "LYU Beisong")
    y -= 14

    c.setFont("Helvetica-Bold", 9.2)
    c.setFillColor(colors.black)
    c.drawCentredString(
        width / 2,
        y,
        "Candidat admissible IMT Nord Europe | Recherche contrat d'apprentissage ingénieur 36 mois | Rentrée 2026",
    )
    y -= 10.5

    c.setFont("Helvetica", 7.8)
    c.setFillColor(colors.HexColor("#555555"))
    c.drawCentredString(
        width / 2,
        y,
        "Tours | Mobilité Lille, Douai et Hauts-de-France | 06 28 73 21 24 | benson201814@gmail.com | LinkedIn / GitHub : à compléter",
    )
    y -= 7
    c.setStrokeColor(colors.HexColor("#BFBFBF"))
    c.line(margin_x, y, width - margin_x, y)

    y = draw_section(c, "Profil", margin_x, y, usable)
    profile = (
        "Admissible à IMT Nord Europe pour une formation d'ingénieur par apprentissage, je recherche un contrat de 36 mois "
        "en informatique, télécommunications, réseaux ou data. Issu d'une CPGE MP2I-MPI, je combine une base solide en "
        "mathématiques, algorithmique et programmation avec un parcours international Chine - États-Unis - France. "
        "Trilingue chinois, anglais et français, je m'adapte rapidement aux environnements exigeants et multiculturels."
    )
    y = draw_wrapped(c, profile, margin_x, y, usable, size=8.25, leading=9.35)

    y = draw_section(c, "Formation", margin_x, y, usable)
    y = draw_role(c, "2024 - 2026 | CPGE MP2I-MPI - Lycée Descartes, Tours", "", margin_x, y, usable)
    y = draw_wrapped(
        c,
        "Mathématiques, informatique, algorithmique, physique. Compétences développées : rigueur scientifique, résolution de problèmes complexes, autonomie, forte capacité de travail. Responsabilités : délégué, membre du conseil d'administration.",
        margin_x + 8,
        y,
        usable - 8,
        size=8.0,
        leading=9.0,
    )
    y = draw_role(c, "2026 - 2029 | IMT Nord Europe", "admissible, contrat d'apprentissage à finaliser", margin_x, y, usable)
    y = draw_wrapped(
        c,
        "Formation d'ingénieur par apprentissage visée. Domaine cible : Informatique, Télécommunications et Réseaux.",
        margin_x + 8,
        y,
        usable - 8,
        size=8.0,
        leading=9.0,
    )
    y = draw_role(c, "2021 - 2024 | Baccalauréat général - Lycée Jean Lurçat", "", margin_x, y, usable)
    y = draw_wrapped(
        c,
        "Spécialités scientifiques. Vice-président du CVL, membre du conseil d'administration, délégué.",
        margin_x + 8,
        y,
        usable - 8,
        size=8.0,
        leading=9.0,
    )

    y = draw_section(c, "Competences techniques", margin_x, y, usable)
    for line in [
        "Programmation : Python, C, OCaml, bases SQL | Informatique : algorithmique, structures de données, programmation",
        "Signal et données : analyse fréquentielle, traitement numérique du signal, modélisation",
        "Outils et méthodes : Git, Linux, Windows | analyse, synthèse, résolution de problèmes, apprentissage rapide",
    ]:
        y = draw_wrapped(c, line, margin_x + 8, y, usable - 8, size=8.0, leading=9.0)

    y = draw_section(c, "Projets", margin_x, y, usable)
    y = draw_role(c, "TIPE - Analyse et synthèse du signal audio", "", margin_x, y, usable)
    y = draw_wrapped(c, "Implémentation d'algorithmes de synthèse sonore ; segmentation et analyse fréquentielle de signaux audio ; présentation d'une démarche scientifique structurée.", margin_x + 8, y, usable - 8, size=7.95, leading=8.8)
    y = draw_role(c, "Développement de jeux indépendants", "", margin_x, y, usable)
    y = draw_wrapped(c, "Conception de prototypes interactifs ; programmation de mécaniques de jeu et de logique applicative ; tests et amélioration en autonomie.", margin_x + 8, y, usable - 8, size=7.95, leading=8.8)
    y = draw_role(c, "Création de contenu audiovisuel", "", margin_x, y, usable)
    y = draw_wrapped(c, "Production de plus de 30 projets audiovisuels, 200 000 vues cumulées. Compétences : narration, montage, régularité et communication.", margin_x + 8, y, usable - 8, size=7.95, leading=8.8)

    y = draw_section(c, "Experiences et engagements", margin_x, y, usable)
    y = draw_role(c, "2023 | Assistant traducteur et support technique", "Projet UiT The Arctic University of Norway", margin_x, y, usable)
    y = draw_wrapped(c, "Création et gestion d'un site web ; traduction et interprétation chinois - anglais - français ; coordination et communication dans un contexte international.", margin_x + 8, y, usable - 8, size=7.95, leading=8.8)
    y = draw_role(c, "2021 | Volontariat d'enseignement de l'anglais", "Tibet, Chine", margin_x, y, usable)
    y = draw_wrapped(c, "Enseignement auprès d'élèves en zone rurale ; animation de groupe et adaptation pédagogique.", margin_x + 8, y, usable - 8, size=7.95, leading=8.8)
    y = draw_role(c, "2020 | Interprète bilingue", "Camp d'été, Portland, États-Unis", margin_x, y, usable)
    y = draw_wrapped(c, "Traduction chinois - anglais ; support logistique et médiation interculturelle.", margin_x + 8, y, usable - 8, size=7.95, leading=8.8)

    y = draw_section(c, "Langues et centres d'interet", margin_x, y, usable)
    y = draw_wrapped(c, "Langues : chinois langue maternelle | anglais courant, 4 ans de scolarité aux États-Unis | français courant", margin_x + 8, y, usable - 8, size=8.0, leading=9.0)
    y = draw_wrapped(c, "Intérêts : piano | course à pied | jeux vidéo indépendants | communication interculturelle", margin_x + 8, y, usable - 8, size=8.0, leading=9.0)

    c.showPage()
    c.save()
    print(PDF_PATH.resolve())


if __name__ == "__main__":
    build_cv()
    build_pdf()
